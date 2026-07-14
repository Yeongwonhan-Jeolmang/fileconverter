"""Document conversion: text, markdown, HTML, DOCX, PDF, RTF.

Layered strategy (most specific -> most general):
  1. Native, dependency-light paths for common pairs (txt/md/html/docx/pdf).
  2. If LibreOffice ("soffice") is found on PATH, use it as a universal
     fallback engine.
  3. If Pandoc is found on PATH, use it as a secondary fallback for
     text-based formats (md/html/rst/docx/odt/etc.).
"""

from __future__ import annotations

import shutil
import subprocess
from pathlib import Path

from ..core.base import BaseConverter, ConversionJob, ConversionResult, ProgressCallback
from ..core.exceptions import ConversionFailedError, MissingDependencyError
from ..core.registry import register

_TEXT_LIKE = frozenset({"txt", "md", "html", "htm"})
_OFFICE_LIKE = frozenset({"docx", "doc", "odt", "rtf", "pdf", "pptx", "ppt"})
_ALL_DOC_FORMATS = _TEXT_LIKE | _OFFICE_LIKE


def _find_libreoffice() -> str | None:
    for candidate in ("soffice", "libreoffice"):
        path = shutil.which(candidate)
        if path:
            return path
    return None


def _find_pandoc() -> str | None:
    return shutil.which("pandoc")


class DocumentConverter(BaseConverter):
    name = "Document Converter"
    description = (
        "Converts between TXT, Markdown, HTML, DOCX, RTF, ODT, PPTX and PDF. "
        "Uses native libraries for common pairs and auto-detects LibreOffice/Pandoc "
        "on your system as fallback engines for everything else."
    )
    input_formats = _ALL_DOC_FORMATS
    output_formats = _ALL_DOC_FORMATS

    def check_available(self) -> tuple[bool, str]:
        notes = []
        if _find_libreoffice():
            notes.append("LibreOffice found (full office format support)")
        else:
            notes.append("LibreOffice not found (install for docx/pdf/odt conversions)")
        if _find_pandoc():
            notes.append("Pandoc found (extra text format support)")
        return True, "; ".join(notes)

    def convert(self, job: ConversionJob, progress_cb: ProgressCallback = None) -> ConversionResult:
        def _do() -> None:
            src_ext = job.source_path.suffix.lower().lstrip(".")
            target = job.target_format.lower().lstrip(".")

            if progress_cb:
                progress_cb(0.1, f"Converting {src_ext} -> {target}")

            if self._try_native(job, src_ext, target, progress_cb):
                return
            if self._try_libreoffice(job, target, progress_cb):
                return
            if self._try_pandoc(job, src_ext, target, progress_cb):
                return

            raise MissingDependencyError(
                "LibreOffice or Pandoc",
                f"No native path exists for {src_ext} -> {target}. Install LibreOffice "
                "(https://libreoffice.org) or Pandoc (https://pandoc.org) for full document support.",
            )

        return self._run_timed(job, _do)

    # ------------------------------------------------------------------
    def _try_native(self, job: ConversionJob, src_ext: str, target: str, progress_cb) -> bool:
        if src_ext in ("txt",) and target in ("md", "html", "htm"):
            text = job.source_path.read_text(encoding="utf-8", errors="replace")
            if target == "md":
                job.output_path.write_text(text, encoding="utf-8")
            else:
                escaped = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                html = f"<html><body><pre>{escaped}</pre></body></html>"
                job.output_path.write_text(html, encoding="utf-8")
            return True

        if src_ext == "md" and target in ("html", "htm"):
            try:
                import markdown  # type: ignore
            except ImportError as exc:
                raise MissingDependencyError("markdown", "Install with: pip install markdown") from exc
            text = job.source_path.read_text(encoding="utf-8", errors="replace")
            html_body = markdown.markdown(text, extensions=["extra", "tables", "fenced_code"])
            html = f"<html><head><meta charset='utf-8'></head><body>{html_body}</body></html>"
            job.output_path.write_text(html, encoding="utf-8")
            return True

        if src_ext in ("html", "htm") and target == "md":
            try:
                import html2text  # type: ignore
            except ImportError as exc:
                raise MissingDependencyError("html2text", "Install with: pip install html2text") from exc
            html = job.source_path.read_text(encoding="utf-8", errors="replace")
            converter = html2text.HTML2Text()
            job.output_path.write_text(converter.handle(html), encoding="utf-8")
            return True

        if src_ext in ("html", "htm", "md") and target == "txt":
            text = job.source_path.read_text(encoding="utf-8", errors="replace")
            if src_ext in ("html", "htm"):
                try:
                    from bs4 import BeautifulSoup  # type: ignore

                    text = BeautifulSoup(text, "html.parser").get_text()
                except ImportError:
                    pass
            job.output_path.write_text(text, encoding="utf-8")
            return True

        if src_ext == "docx" and target == "txt":
            try:
                import docx  # type: ignore
            except ImportError as exc:
                raise MissingDependencyError("python-docx", "Install with: pip install python-docx") from exc
            document = docx.Document(str(job.source_path))
            text = "\n".join(p.text for p in document.paragraphs)
            job.output_path.write_text(text, encoding="utf-8")
            return True

        if src_ext == "docx" and target in ("html", "htm"):
            try:
                import docx  # type: ignore
            except ImportError as exc:
                raise MissingDependencyError("python-docx", "Install with: pip install python-docx") from exc
            document = docx.Document(str(job.source_path))
            parts = [f"<p>{p.text}</p>" for p in document.paragraphs]
            html = f"<html><body>{''.join(parts)}</body></html>"
            job.output_path.write_text(html, encoding="utf-8")
            return True

        if src_ext in ("txt", "md") and target == "pdf":
            self._text_to_pdf(job)
            return True

        if src_ext == "pdf" and target == "txt":
            try:
                from pypdf import PdfReader  # type: ignore
            except ImportError as exc:
                raise MissingDependencyError("pypdf", "Install with: pip install pypdf") from exc
            reader = PdfReader(str(job.source_path))
            text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
            job.output_path.write_text(text, encoding="utf-8")
            return True

        return False

    def _text_to_pdf(self, job: ConversionJob) -> None:
        try:
            from reportlab.lib.pagesizes import LETTER  # type: ignore
            from reportlab.pdfgen import canvas
        except ImportError as exc:
            raise MissingDependencyError("reportlab", "Install with: pip install reportlab") from exc

        text = job.source_path.read_text(encoding="utf-8", errors="replace")
        c = canvas.Canvas(str(job.output_path), pagesize=LETTER)
        width, height = LETTER
        y = height - 72
        for line in text.splitlines() or [""]:
            for wrapped in _wrap_line(line, 95):
                if y < 72:
                    c.showPage()
                    y = height - 72
                c.drawString(72, y, wrapped)
                y -= 14
        c.save()

    def _try_libreoffice(self, job: ConversionJob, target: str, progress_cb) -> bool:
        soffice = _find_libreoffice()
        if not soffice:
            return False
        out_dir = job.output_path.parent
        cmd = [
            soffice,
            "--headless",
            "--norestore",
            "--convert-to",
            target,
            "--outdir",
            str(out_dir),
            str(job.source_path),
        ]
        if progress_cb:
            progress_cb(0.4, "Running LibreOffice headless conversion")
        proc = subprocess.run(cmd, capture_output=True, text=True, timeout=180)
        if proc.returncode != 0:
            raise ConversionFailedError(f"LibreOffice conversion failed: {proc.stderr.strip()}")
        produced = out_dir / f"{job.source_path.stem}.{target}"
        if produced.exists() and produced != job.output_path:
            produced.replace(job.output_path)
        return job.output_path.exists()

    def _try_pandoc(self, job: ConversionJob, src_ext: str, target: str, progress_cb) -> bool:
        pandoc = _find_pandoc()
        if not pandoc:
            return False
        if progress_cb:
            progress_cb(0.4, "Running Pandoc")
        cmd = [pandoc, str(job.source_path), "-o", str(job.output_path)]
        proc = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
        if proc.returncode != 0:
            raise ConversionFailedError(f"Pandoc conversion failed: {proc.stderr.strip()}")
        return job.output_path.exists()


def _wrap_line(line: str, width: int) -> list[str]:
    if not line:
        return [""]
    return [line[i : i + width] for i in range(0, len(line), width)]


register(DocumentConverter())
